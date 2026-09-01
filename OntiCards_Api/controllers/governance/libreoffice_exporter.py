# -*- coding:utf-8 -*-
"""
@File: libreoffice_exporter.py
@Description: 数据治理报告文档导出服务 - 基于已有报告数据生成 Word/PDF/Excel 文档
@Author: 韩小豪 849631113@qq.com
@Create: 2026-06-29
@Update: 2026-07-07

文档生成优先级（两级，降级策略）：
1. 本地 soffice（优先）：容器内已安装 LibreOffice，通过 subprocess 调用 soffice 转换
   - 依赖：Dockerfile 中安装 libreoffice 包
   - 支持：DOCX / PDF / XLSX / ODT / ODS
2. 纯 Python 方案（降级）：
   - DOCX: python-docx
   - XLSX: openpyxl
   - PDF: 暂不支持（soffice 不可用时直接报错）

导出格式：
- docx: Word 2007+ 文档
- pdf: PDF 文档（通过 soffice 转换 HTML 生成）
- xlsx: Excel 2007+ 文档

已移除：
- weasyprint PDF 降级方案（依赖 GTK，在 Windows/Linux 环境配置复杂，且效果不如 soffice）
- LibreOffice HTTP 容器服务调用（LIBREOFFICE_SERVICE_URL 等配置已废弃）
"""

import os
import re
import subprocess
import tempfile
import shutil
import json
from datetime import datetime
from typing import Optional, Dict, Any, List

from models.governance_report import GovernanceReport
from models.rule_execution_result import RuleExecutionResult


# ==================== 格式枚举 ====================

class ExportFormat:
    DOCX = 'docx'
    PDF = 'pdf'
    XLSX = 'xlsx'
    HTML = 'html'
    ODT = 'odt'
    ODS = 'ods'

    SUPPORTED = [DOCX, PDF, XLSX]
    LIBREOFFICE_FORMATS = [DOCX, PDF, ODT, ODS]


# ==================== 工具函数（soffice subprocess 调用） ====================

def _get_local_soffice_path() -> Optional[str]:
    """查找本地 soffice 可执行文件路径（通过 PATH 环境变量）

    适用于 soffice 与 API 进程在同一台机器上的场景，包括 Docker 容器内安装的 LibreOffice。

    Windows: 优先 soffice.exe，过滤 .com；Linux/Docker: 直接取 which 结果。
    """
    candidates = []
    try:
        if os.name == 'nt':
            result = subprocess.run(
                ['where', 'soffice'],
                capture_output=True, text=True, timeout=5
            )
            if result.returncode == 0:
                candidates = result.stdout.strip().split('\n')
        else:
            result = subprocess.run(
                ['which', 'soffice'],
                capture_output=True, text=True, timeout=5
            )
            if result.returncode == 0:
                candidates = result.stdout.strip().split('\n')
    except Exception:
        pass

    # Windows: 优先 soffice.exe，过滤 .com（.com 对 filter 系统支持不完整）
    if os.name == 'nt':
        exe_paths = [p for p in candidates if p.lower().endswith('.exe') and not p.lower().endswith('.com')]
        com_paths = [p for p in candidates if p.lower().endswith('.com')]
        if exe_paths:
            return exe_paths[0]
        if com_paths:
            return com_paths[0]
        # 兜底：常见 Windows 安装路径
        for base in [
            r"C:\Program Files\LibreOffice\program",
            r"C:\Program Files (x86)\LibreOffice\program",
            r"H:\LibreOffice\program",
        ]:
            exe = os.path.join(base, "soffice.exe")
            if os.path.isfile(exe):
                return exe
    else:
        # Linux/Docker: 直接取第一个
        if candidates:
            return candidates[0]
    return None


# 根据目标格式选择对应的输入 filter name（用于修复 Windows 上 LibreOffice 25.x
# export filter 查找失败的问题）。
# 当源文件为 HTML 时需要明确指定输入 filter，否则 soffice 无法正确定位 export filter。
_FILTER_MAP = {
    'docx': 'HTML (StarWriter)',
    'xlsx': 'HTML (StarCalc)',
}


def _run_local_soffice(src_path: str, dst_dir: str, target_format: str) -> Optional[str]:
    """调用本地 soffice 转换文件（优先方案）

    Args:
        src_path: 源文件路径（支持 HTML, ODT 等）
        dst_dir: 输出目录
        target_format: 目标格式（docx, pdf, xlsx 等）

    Returns:
        输出文件路径，失败返回 None
    """
    soffice = _get_local_soffice_path()
    if not soffice:
        return None

    try:
        cmd = [
            soffice,
            '--headless',
            '--convert-to', target_format,
        ]
        infilter = _FILTER_MAP.get(target_format.lower())
        if infilter:
            cmd.append('--infilter=' + infilter)
        cmd.extend(['--outdir', dst_dir, src_path])

        env = os.environ.copy()
        if os.name == 'nt':
            env['HOME'] = tempfile.gettempdir()
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,
            env=env,
        )
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(src_path))[0]
            output_name = f"{base_name}.{target_format}"
            output_path = os.path.join(dst_dir, output_name)
            if os.path.isfile(output_path):
                return output_path
            for f in os.listdir(dst_dir):
                if f.startswith(base_name) and f.endswith(f'.{target_format}'):
                    return os.path.join(dst_dir, f)
        return None
    except Exception as e:
        print(f"[WARN] Local soffice conversion failed: {e}")
        return None


# ==================== HTML 模板引擎 ====================

def _build_report_html(report: GovernanceReport,
                       results: List[RuleExecutionResult]) -> str:
    """生成报告 HTML 内容（用于 LibreOffice 转换）"""
    total = len(results)
    passed = sum(1 for r in results if r.status == 'passed')
    failed = sum(1 for r in results if r.status == 'failed')
    errors = sum(1 for r in results if r.status == 'error')
    quality_score = float(report.quality_score) if report.quality_score else 0
    grade = report.grade or '一般'
    failed_total_rows = sum(r.failed_count or 0 for r in results if r.status == 'failed')

    # 提取关系盘点数据（来自 report.details）
    relation_discovery = None
    if report.details and isinstance(report.details, dict):
        relation_discovery = report.details.get('relation_discovery')

    grade_colors = {
        '优秀': '#27ae60',
        '良好': '#3498db',
        '一般': '#f39c12',
        '较差': '#e67e22',
        '差': '#e74c3c'
    }
    grade_color = grade_colors.get(grade, '#95a5a6')
    score_color = grade_color

    rule_type_names = {
        'null_check': '空值检测',
        'unique': '唯一性检测',
        'format': '格式检测',
        'threshold': '阈值检测',
        'enum': '枚举检测',
        'custom_sql': '自定义SQL',
        'length_check': '长度检测',
        'range_check': '范围检测',
        'date_check': '日期检测',
        'consistency_check': '一致性检测',
        'freshness_check': '新鲜度检测',
        'value_distribution': '值分布检测',
        'composite': '复合条件检测',
        'table_stats': '表级统计',
    }
    severity_names = {
        'critical': '严重',
        'warning': '警告',
        'info': '信息'
    }
    status_names = {
        'passed': '通过',
        'failed': '失败',
        'error': '错误'
    }

    results_rows = ''
    for r in results:
        status_txt = status_names.get(r.status, r.status or 'N/A')
        severity_txt = severity_names.get(r.severity, r.severity or 'N/A')
        rule_type_txt = rule_type_names.get(r.rule_type, r.rule_type or 'N/A')
        failed_rate_txt = f"{float(r.failed_rate):.2f}%" if r.failed_rate is not None else 'N/A'
        failed_count_txt = str(r.failed_count) if r.failed_count is not None else '0'
        passed_count_txt = str(r.passed_count) if r.passed_count is not None else '0'
        total_count_txt = str(r.total_count) if r.total_count is not None else '0'
        status_color = '#27ae60' if r.status == 'passed' else ('#e74c3c' if r.status == 'failed' else '#f39c12')

        results_rows += f'''
        <tr>
            <td style="padding:8px;border:1px solid #ddd;">{r.rule_name or "N/A"}</td>
            <td style="padding:8px;border:1px solid #ddd;">{r.table_name or "N/A"}</td>
            <td style="padding:8px;border:1px solid #ddd;">{r.column_name or "N/A"}</td>
            <td style="padding:8px;border:1px solid #ddd;">{rule_type_txt}</td>
            <td style="padding:8px;border:1px solid #ddd;">{severity_txt}</td>
            <td style="padding:8px;border:1px solid #ddd;color:{status_color};font-weight:bold;">{status_txt}</td>
            <td style="padding:8px;border:1px solid #ddd;text-align:right;">{failed_rate_txt}</td>
            <td style="padding:8px;border:1px solid #ddd;text-align:right;">{failed_count_txt}/{total_count_txt}</td>
        </tr>
        '''

    # 构建关系发现章节（仅当有数据时渲染）
    relation_section = ''
    if relation_discovery and relation_discovery.get('relationships_count', 0) > 0:
        rel_count = relation_discovery.get('relationships_count', 0)
        tables_count = relation_discovery.get('tables_count', 0)
        cards_count = relation_discovery.get('cards_count', 0)
        cross_source_count = relation_discovery.get('cross_source_count', 0)
        is_multi_source = relation_discovery.get('is_multi_source', False)
        statistics = relation_discovery.get('statistics', {}) or {}

        rel_summary_cards = f'''
    <div class="summary-card">
        <div class="summary-value">{rel_count}</div>
        <div class="summary-label">发现关系数</div>
    </div>
    <div class="summary-card">
        <div class="summary-value">{tables_count}</div>
        <div class="summary-label">涉及表数</div>
    </div>
    <div class="summary-card">
        <div class="summary-value">{cards_count}</div>
        <div class="summary-label">生成卡片数</div>
    </div>
    <div class="summary-card">
        <div class="summary-value">{cross_source_count}</div>
        <div class="summary-label">跨源关系数</div>
    </div>'''

        rel_source_tag = '<span style="background:#e8f5e9;color:#2e7d32;padding:2px 8px;border-radius:4px;font-size:12px;">多数据源</span>' if is_multi_source else '<span style="background:#f3f4f6;color:#666;padding:2px 8px;border-radius:4px;font-size:12px;">单数据源</span>'

        stat_rows = ''
        for key, value in statistics.items():
            if isinstance(value, dict):
                stat_rows += f'<tr><td style="padding:6px 12px;border:1px solid #ddd;">{key}</td><td style="padding:6px 12px;border:1px solid #ddd;">{value.get("count", 0)} 条</td></tr>'
            else:
                stat_rows += f'<tr><td style="padding:6px 12px;border:1px solid #ddd;">{key}</td><td style="padding:6px 12px;border:1px solid #ddd;">{value}</td></tr>'

        stat_table_html = f'''
<table style="width:auto;margin:10px 0;">
    <thead>
        <tr><th style="padding:6px 12px;background:#f8f9fa;">指标</th><th style="padding:6px 12px;background:#f8f9fa;">数值</th></tr>
    </thead>
    <tbody>{stat_rows}</tbody>
</table>''' if stat_rows else ''

        relation_section = f'''
<h2 style="margin-top:30px;">三、表关系发现结果 {rel_source_tag}</h2>
<div class="summary-grid">{rel_summary_cards}
</div>
{stat_table_html}'''

    failed_samples_section = ''
    failed_with_samples = [r for r in results if r.status == 'failed' and r.failed_samples]
    if failed_with_samples:
        failed_samples_section = '<h2 style="margin-top:30px;">四、失败样本明细</h2>'
        for r in failed_with_samples[:10]:
            samples_json = r.failed_samples
            if isinstance(samples_json, str):
                try:
                    samples_json = json.loads(samples_json)
                except Exception:
                    samples_json = []
            elif not isinstance(samples_json, list):
                samples_json = []

            violated_cond = ''
            cond_expr = ''
            if samples_json:
                first_sample = samples_json[0]
                violated_cond = first_sample.get('violated_condition', '')
                cond_expr = first_sample.get('condition_expr', '')

            # 避免在 f-string 表达式中使用 \"，改用 .format 避免多行拼接歧义
            cond_line = (
                "<br/>违反条件: <code>" + (violated_cond or "") + "</code>" +
                ("<br/>原始条件: <code>" + (cond_expr or "") + "</code>" if cond_expr else "")
            ) if violated_cond or cond_expr else ""
            sample_display = (
                "<div style='margin-top:6px;color:#555;'>样本值: " +
                str(samples_json[:3]) + "</div>"
            ) if samples_json else ""
            rule_html = (
                "<div style='margin:10px 0;padding:12px;border-left:3px solid #e74c3c;background:#fef9f9;'>"
                "<strong style='color:#e74c3c;'>{}</strong>"
                "<span style='color:#666;'>（{}.{}）</span>"
                "<div style='margin-top:6px;'>"
                "失败率: <strong>{:.2f}%</strong>"
                "失败记录: <strong>{}</strong> 条{}"
                "</div>{}"
                "</div>"
            ).format(
                r.rule_name or "N/A",
                r.table_name or "",
                r.column_name or "",
                float(r.failed_rate or 0),
                r.failed_count or 0,
                cond_line,
                sample_display,
            )
            failed_samples_section += rule_html

    recommendations = _generate_recommendations(results, rule_type_names)

    execution_time = report.execution_time.strftime('%Y-%m-%d %H:%M:%S') if report.execution_time else datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    report_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    datasource_id_str = str(report.datasource_id) if report.datasource_id else 'N/A'

    return f'''<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>{report.report_name or "数据治理质量报告"}</title>
<style>
body {{ font-family: "Microsoft YaHei", "SimHei", Arial, sans-serif; margin: 40px; color: #333; font-size: 14px; }}
h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; margin-top: 0; }}
h2 {{ color: #34495e; margin-top: 30px; border-left: 4px solid #3498db; padding-left: 10px; }}
.header {{ text-align: center; margin-bottom: 30px; }}
.summary-grid {{ display: flex; flex-wrap: wrap; gap: 15px; margin: 20px 0; }}
.summary-card {{
    flex: 1; min-width: 100px; text-align: center;
    padding: 20px 10px; border-radius: 8px;
    background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
    border: 1px solid #dee2e6;
}}
.summary-value {{ font-size: 28px; font-weight: bold; color: #2c3e50; }}
.summary-label {{ font-size: 13px; color: #666; margin-top: 5px; }}
table {{ width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 13px; }}
th {{ background-color: #3498db; color: white; padding: 10px 8px; text-align: left; font-weight: normal; }}
td {{ padding: 8px; border: 1px solid #ddd; }}
tr:nth-child(even) {{ background-color: #f8f9fa; }}
tr:hover {{ background-color: #e8f4fd; }}
.recommendation {{ margin: 8px 0; padding: 8px 12px; border-left: 3px solid #3498db; background: #f0f8ff; }}
.footer {{ margin-top: 40px; text-align: center; color: #999; font-size: 12px; border-top: 1px solid #ddd; padding-top: 15px; }}
.info-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin: 15px 0; }}
.info-item {{ padding: 8px 12px; background: #f8f9fa; border-radius: 4px; }}
.info-label {{ color: #666; font-size: 12px; }}
.info-value {{ font-weight: bold; margin-top: 2px; }}
</style>
</head>
<body>

<div class="header">
    <h1>数据治理质量报告</h1>
    <p style="font-size:16px;color:#666;">{report.report_name or "数据质量治理盘点报告"}</p>
</div>

<h2>一、基本信息</h2>
<div class="info-grid">
    <div class="info-item">
        <div class="info-label">数据源</div>
        <div class="info-value">{datasource_id_str}</div>
    </div>
    <div class="info-item">
        <div class="info-label">执行时间</div>
        <div class="info-value">{execution_time}</div>
    </div>
    <div class="info-item">
        <div class="info-label">生成时间</div>
        <div class="info-value">{report_time}</div>
    </div>
    <div class="info-item">
        <div class="info-label">涉及表</div>
        <div class="info-value">{", ".join(report.scope_tables) if report.scope_tables else "全部"}</div>
    </div>
</div>

<h2>二、质量概览</h2>
<div class="summary-grid">
    <div class="summary-card" style="border-left:4px solid {score_color};">
        <div class="summary-value" style="color:{score_color};">{quality_score:.1f}</div>
        <div class="summary-label">质量评分</div>
    </div>
    <div class="summary-card" style="border-left:4px solid {grade_color};">
        <div class="summary-value" style="color:{grade_color};">{grade}</div>
        <div class="summary-label">质量评级</div>
    </div>
    <div class="summary-card">
        <div class="summary-value">{total}</div>
        <div class="summary-label">执行规则数</div>
    </div>
    <div class="summary-card" style="border-left:4px solid #27ae60;">
        <div class="summary-value" style="color:#27ae60;">{passed}</div>
        <div class="summary-label">通过</div>
    </div>
    <div class="summary-card" style="border-left:4px solid #e74c3c;">
        <div class="summary-value" style="color:#e74c3c;">{failed}</div>
        <div class="summary-label">失败</div>
    </div>
    <div class="summary-card">
        <div class="summary-value">{failed_total_rows}</div>
        <div class="summary-label">影响行数</div>
    </div>
</div>

{relation_section}

<h2>四、执行明细</h2>
<table>
    <thead>
        <tr>
            <th>规则名称</th>
            <th>目标表</th>
            <th>目标列</th>
            <th>规则类型</th>
            <th>严重级别</th>
            <th>状态</th>
            <th>失败率</th>
            <th>失败/总数</th>
        </tr>
    </thead>
    <tbody>
        {results_rows}
    </tbody>
</table>

{failed_samples_section}

<h2>五、改进建议</h2>
{recommendations}

<div class="footer">
    <p>本报告由 OntiCards AI数据中枢 自动生成</p>
    <p>报告生成时间: {report_time}</p>
</div>

</body>
</html>'''


def _generate_recommendations(results: List[RuleExecutionResult],
                               rule_type_names: dict) -> str:
    """生成改进建议 HTML"""
    failed_results = [r for r in results if r.status == 'failed']
    if not failed_results:
        return '<p style="color:#27ae60;">&#10004; 数据质量良好，未发现问题，建议继续保持当前的数据管理规范。</p>'

    by_type = {}
    for r in failed_results:
        rt = r.rule_type or 'unknown'
        if rt not in by_type:
            by_type[rt] = []
        by_type[rt].append(r)

    recommendations = []
    priority = [
        ('null_check', '空值问题'),
        ('unique', '唯一性问题'),
        ('format', '格式问题'),
        ('threshold', '阈值问题'),
        ('enum', '枚举问题'),
        ('custom_sql', '自定义规则问题'),
    ]

    for rule_type, label in priority:
        if rule_type in by_type:
            items = by_type[rule_type]
            tables = list(set(r.table_name for r in items if r.table_name))
            total_rows = sum(r.failed_count or 0 for r in items)
            tbl_list = ', '.join(tables[:3])
            if len(tables) > 3:
                tbl_list += f' 等{len(tables)}个表'

            advice_map = {
                'null_check': f'发现 {len(items)} 个字段存在空值，共 {total_rows} 条记录，建议检查 {tbl_list} 等表的数据录入流程，确保必填字段有值。',
                'unique': f'发现 {len(items)} 个字段存在重复值，共 {total_rows} 条重复记录，建议对 {tbl_list} 等表进行去重处理。',
                'format': f'发现 {len(items)} 个字段存在格式不规范问题，建议统一 {tbl_list} 等表的字段录入格式标准。',
                'threshold': f'发现 {len(items)} 个字段存在超出阈值范围的数据，建议检查 {tbl_list} 等表数据来源的合法性。',
                'enum': f'发现 {len(items)} 个字段存在非法枚举值，建议规范 {tbl_list} 等表的数据取值范围。',
                'custom_sql': f'发现 {len(items)} 条自定义规则检测失败，建议检查对应的 SQL 条件是否需要调整。',
            }
            recommendations.append(f'<div class="recommendation"><strong>{label}：</strong>{advice_map.get(rule_type, "建议进行检查和修复。")}</div>')

    for rt, items in by_type.items():
        if rt not in dict(priority):
            tables = list(set(r.table_name for r in items if r.table_name))
            recommendations.append(f'<div class="recommendation"><strong>{rule_type_names.get(rt, rt)}：</strong>发现 {len(items)} 处问题，涉及表: {", ".join(tables[:3])} 等，建议按类型针对性修复。</div>')

    return '\n'.join(recommendations) if recommendations else '<p>建议继续保持当前的数据管理规范。</p>'


# ==================== 纯 Python 生成器（不依赖 LibreOffice） ====================

def _build_docx_with_python(report: GovernanceReport,
                             results: List[RuleExecutionResult],
                             output_path: str) -> bool:
    """使用 python-docx 生成 Word 文档（不依赖 LibreOffice）"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    doc = Document()
    title = doc.add_heading(report.report_name or '数据治理质量报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'执行时间: {report.execution_time.strftime("%Y-%m-%d %H:%M:%S") if report.execution_time else "N/A"}')
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'数据源: {report.datasource_id}')

    doc.add_heading('质量概览', level=1)
    score = float(report.quality_score) if report.quality_score else 0
    grade = report.grade or '一般'
    doc.add_paragraph(f'质量评分: {score:.1f}  分    评级: {grade}')
    doc.add_paragraph(f'执行规则: {len(results)} 条    通过: {sum(1 for r in results if r.status=="passed")}    失败: {sum(1 for r in results if r.status=="failed")}')

    # 关系发现章节
    relation_discovery = None
    if report.details and isinstance(report.details, dict):
        relation_discovery = report.details.get('relation_discovery')
    if relation_discovery and relation_discovery.get('relationships_count', 0) > 0:
        doc.add_heading('表关系发现结果', level=1)
        doc.add_paragraph(f"发现关系数: {relation_discovery.get('relationships_count', 0)}    "
                          f"涉及表数: {relation_discovery.get('tables_count', 0)}    "
                          f"生成卡片数: {relation_discovery.get('cards_count', 0)}    "
                          f"跨源关系数: {relation_discovery.get('cross_source_count', 0)}")
        is_multi = relation_discovery.get('is_multi_source', False)
        doc.add_paragraph(f"数据源类型: {'多数据源' if is_multi else '单数据源'}")
        stats = relation_discovery.get('statistics') or {}
        if stats:
            doc.add_paragraph("统计信息:")
            for k, v in stats.items():
                val_str = f"{v.get('count', 0)} 条" if isinstance(v, dict) else str(v)
                doc.add_paragraph(f"  {k}: {val_str}")

    doc.add_heading('执行明细', level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['规则名称', '目标表', '目标列', '规则类型', '严重级别', '状态', '失败率', '失败/总数']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    rule_type_names = {
        'null_check': '空值检测', 'unique': '唯一性检测',
        'format': '格式检测', 'threshold': '阈值检测',
        'enum': '枚举检测', 'custom_sql': '自定义SQL',
        'length_check': '长度检测', 'range_check': '范围检测',
        'date_check': '日期检测', 'consistency_check': '一致性检测',
        'freshness_check': '新鲜度检测', 'value_distribution': '值分布检测',
    }
    severity_names = {'critical': '严重', 'warning': '警告', 'info': '信息'}
    status_names = {'passed': '通过', 'failed': '失败', 'error': '错误'}

    for r in results:
        row_cells = table.add_row().cells
        row_cells[0].text = r.rule_name or 'N/A'
        row_cells[1].text = r.table_name or 'N/A'
        row_cells[2].text = r.column_name or 'N/A'
        row_cells[3].text = rule_type_names.get(r.rule_type, r.rule_type or 'N/A')
        row_cells[4].text = severity_names.get(r.severity, r.severity or 'N/A')
        row_cells[5].text = status_names.get(r.status, r.status or 'N/A')
        row_cells[6].text = f"{float(r.failed_rate or 0):.2f}%" if r.failed_rate is not None else 'N/A'
        row_cells[7].text = f"{r.failed_count or 0}/{r.total_count or 0}"

    failed_results = [r for r in results if r.status == 'failed' and r.failed_samples]
    if failed_results:
        doc.add_heading('失败样本', level=1)
        for r in failed_results[:10]:
            doc.add_paragraph(f"规则: {r.rule_name}  表:{r.table_name}  列:{r.column_name}  失败率:{float(r.failed_rate or 0):.2f}%  失败数:{r.failed_count}")
            if r.failed_samples:
                samples = r.failed_samples
                if isinstance(samples, str):
                    try:
                        samples = json.loads(samples)
                    except Exception:
                        samples = []
                for s in samples[:3]:
                    doc.add_paragraph(f"  样本: {s}", style='Intense Quote')

    doc.save(output_path)
    return True


def _build_xlsx_with_python(report: GovernanceReport,
                              results: List[RuleExecutionResult],
                              output_path: str) -> bool:
    """使用 openpyxl 生成 Excel 文档"""
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment
    except ImportError:
        raise ImportError("请安装 openpyxl: pip install openpyxl")

    wb = openpyxl.Workbook()

    ws_summary = wb.active
    ws_summary.title = '概览'
    score = float(report.quality_score) if report.quality_score else 0
    grade = report.grade or '一般'
    overview_data = [
        ('报告名称', report.report_name or 'N/A'),
        ('执行时间', report.execution_time.strftime('%Y-%m-%d %H:%M:%S') if report.execution_time else 'N/A'),
        ('生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('数据源', str(report.datasource_id) if report.datasource_id else 'N/A'),
        ('质量评分', score),
        ('评级', grade),
        ('执行规则数', len(results)),
        ('通过数', sum(1 for r in results if r.status == 'passed')),
        ('失败数', sum(1 for r in results if r.status == 'failed')),
        ('错误数', sum(1 for r in results if r.status == 'error')),
    ]
    for row_idx, (k, v) in enumerate(overview_data, 1):
        ws_summary.cell(row=row_idx, column=1, value=k)
        ws_summary.cell(row=row_idx, column=2, value=v)

    # 关系发现章节
    relation_discovery = None
    if report.details and isinstance(report.details, dict):
        relation_discovery = report.details.get('relation_discovery')
    if relation_discovery and relation_discovery.get('relationships_count', 0) > 0:
        ws_rel = wb.create_sheet('关系发现结果')
        ws_rel.append(['指标', '数值'])
        ws_rel.append(['发现关系数', relation_discovery.get('relationships_count', 0)])
        ws_rel.append(['涉及表数', relation_discovery.get('tables_count', 0)])
        ws_rel.append(['生成卡片数', relation_discovery.get('cards_count', 0)])
        ws_rel.append(['跨源关系数', relation_discovery.get('cross_source_count', 0)])
        ws_rel.append(['数据源类型', '多数据源' if relation_discovery.get('is_multi_source') else '单数据源'])
        stats = relation_discovery.get('statistics') or {}
        for k, v in stats.items():
            val_str = f"{v.get('count', 0)} 条" if isinstance(v, dict) else str(v)
            ws_rel.append([k, val_str])

    ws_results = wb.create_sheet('执行明细')
    headers = ['规则名称', '规则类型', '严重级别', '目标表', '目标列',
               '状态', '总记录数', '通过数', '失败数', '失败率(%)',
               'library_id', 'rule_id']
    ws_results.append(headers)
    for r in results:
        ws_results.append([
            r.rule_name or 'N/A',
            r.rule_type or 'N/A',
            r.severity or 'N/A',
            r.table_name or 'N/A',
            r.column_name or 'N/A',
            r.status or 'N/A',
            r.total_count or 0,
            r.passed_count or 0,
            r.failed_count or 0,
            float(r.failed_rate or 0) if r.failed_rate is not None else 0,
            str(r.library_id) if r.library_id else '',
            str(r.rule_id) if r.rule_id else '',
        ])

    ws_samples = wb.create_sheet('失败样本')
    ws_samples.append(['规则名称', '目标表', '目标列', '状态', '失败率(%)', '失败数', '失败样本'])
    for r in results:
        if r.status == 'failed' and r.failed_samples:
            samples = r.failed_samples
            if isinstance(samples, str):
                try:
                    samples = json.loads(samples)
                except Exception:
                    samples = []
            samples_str = json.dumps(samples[:5], ensure_ascii=False)
            ws_samples.append([
                r.rule_name or 'N/A',
                r.table_name or 'N/A',
                r.column_name or 'N/A',
                r.status or 'N/A',
                float(r.failed_rate or 0) if r.failed_rate is not None else 0,
                r.failed_count or 0,
                samples_str,
            ])

    wb.save(output_path)
    return True


# ==================== 主导出类 ====================

class LibreOfficeExporter:
    """数据治理报告文档导出服务

    文档生成优先级（两级，降级策略）：
    1. 本地 soffice（优先）：容器内已安装 LibreOffice，通过 subprocess 调用 soffice 转换
       - 支持：DOCX / PDF / XLSX / ODT / ODS
       - 依赖：Dockerfile 中安装 libreoffice 包
    2. 纯 Python 方案（降级）：
       - DOCX: python-docx
       - XLSX: openpyxl
       - PDF: 暂不支持（soffice 不可用时直接报错）
    """

    def __init__(self, export_path: str = None, user_id: str = None):
        if export_path is None:
            from flask import current_app
            base_path = current_app.config.get(
                'GOVERNANCE_EXPORT_PATH',
                os.path.join(
                    os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
                    'exports', 'governance'
                )
            )
        else:
            base_path = export_path

        if user_id:
            export_path = os.path.join(base_path, str(user_id))
        else:
            export_path = base_path

        # 统一使用正斜杠 '/'，确保跨平台一致性
        self.export_path = export_path.replace('\\', '/')
        self.user_id = user_id
        os.makedirs(self.export_path, exist_ok=True)
        print(f"[LibreOffice-Exporter] 初始化 | export_path={self.export_path}")

    def _check_soffice_available(self) -> Optional[str]:
        """检测 soffice 是否可用，返回路径，找不到返回 None"""
        return _get_local_soffice_path()

    def _make_filename(self, report: GovernanceReport, fmt: str, custom_file_name: str = None) -> str:
        """生成带时间戳的文件名

        Args:
            report: 报告对象
            fmt: 文件格式
            custom_file_name: 自定义文件名（不含扩展名），不传则使用默认命名规则
        """
        # 过滤文件名中的非法字符：/ \ : * ? " < > |
        if custom_file_name and custom_file_name.strip():
            # 自定义文件名：添加中文时间戳
            ts = datetime.now().strftime('%Y年%m月%d日%H点%M分%S秒')
            safe_name = re.sub(r'[\\/:*?"<>|]', '_', custom_file_name.strip())
            return f"{safe_name}_{ts}.{fmt}"
        else:
            # 默认文件名：将 report.report_name 中的冒号替换，并转成中文时间格式
            raw_name = report.report_name or 'governance_report'
            safe_name = self._sanitize_and_format_name(raw_name)
            return f"{safe_name}.{fmt}"

    def _sanitize_and_format_name(self, name: str) -> str:
        """清理文件名中的非法字符，并将时间戳转成中文格式

        例如：
        - 质检结果_2026-07-20_14:07:43 -> 质检结果_2026年07月20日14点07分43秒
        - 质检结果_2026-07-20_14_07_43 -> 质检结果_2026年07月20日14点07分43秒
        """
        # 过滤非法字符：/ \ : * ? " < > |
        safe = re.sub(r'[\\/:*?"<>|]', '', name)
        # 匹配时间戳模式：YYYY-MM-DD_HH:MM:SS 或 YYYY-MM-DD_HH_MM_SS
        time_pattern = r'(\d{4})-(\d{2})-(\d{2})[_-](\d{2})[_:](\d{2})[_:](\d{2})'
        def replace_time(m):
            year, month, day, hour, minute, second = m.groups()
            return f"{year}年{month}月{day}日{hour}点{minute}分{second}秒"
        return re.sub(time_pattern, replace_time, safe)

    def export(self,
               report: GovernanceReport,
               results: List[RuleExecutionResult],
               fmt: str = ExportFormat.DOCX,
               custom_file_name: str = None) -> Dict[str, Any]:
        """导出报告文档

        Args:
            report: 报告对象
            results: 规则执行结果列表
            fmt: 导出格式（docx, pdf, xlsx）
            custom_file_name: 自定义导出文件名（不含扩展名），不传则使用默认命名规则

        Returns:
            {
                'file_path': str,      # 文件绝对路径
                'file_name': str,      # 显示文件名
                'file_size': int,      # 文件大小（字节）
                'format': str,          # 格式
                'mode': str,            # 生成模式: 'soffice' | 'python-docx' | 'openpyxl' | 'weasyprint'
            }
        """
        if fmt == ExportFormat.XLSX:
            result = self._export_xlsx(report, results, custom_file_name=custom_file_name)
        elif fmt == ExportFormat.PDF:
            result = self._export_pdf(report, results, custom_file_name=custom_file_name)
        else:
            result = self._export_docx(report, results, custom_file_name=custom_file_name)

        print(f"[LibreOffice-Exporter] 导出完成 | mode={result['mode']}, file_path={result['file_path']}\n")
        return result

    def _convert_via_local_soffice(
        self,
        tmp_html: str,
        target_fmt: str,
        file_name: str,
    ) -> Optional[Dict[str, Any]]:
        """通过本地 soffice 转换文件（优先方案）

        Args:
            tmp_html: 源 HTML 文件路径
            target_fmt: 目标格式（docx, pdf, xlsx）
            file_name: 最终输出文件名

        Returns:
            成功返回结果字典，失败返回 None
        """
        tmp_dir = tempfile.mkdtemp(prefix='onticards_lo_')
        try:
            output_path = _run_local_soffice(tmp_html, tmp_dir, target_fmt)
            if not output_path:
                return None
            final_path = os.path.join(self.export_path, file_name).replace('\\', '/')
            shutil.move(output_path, final_path)
            return {
                'file_path': final_path,
                'file_name': file_name,
                'file_size': os.path.getsize(final_path),
                'format': target_fmt,
                'mode': 'soffice',
            }
        finally:
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass

    def _cleanup_tmp_dir(self, tmp_dir: str):
        """清理临时目录"""
        try:
            shutil.rmtree(tmp_dir)
        except Exception:
            pass

    def _build_and_convert_html(
        self,
        report: GovernanceReport,
        results: List[RuleExecutionResult],
        target_fmt: str,
        custom_file_name: str = None,
    ) -> Dict[str, Any]:
        """生成 HTML 并转换为目标格式

        Args:
            report: 报告对象
            results: 规则执行结果列表
            target_fmt: 目标格式（docx, pdf）
            custom_file_name: 自定义文件名（不含扩展名）
        """
        print(f"\n[LibreOffice-Exporter] 构建 HTML 报告内容...")
        file_name = self._make_filename(report, target_fmt, custom_file_name=custom_file_name)

        # 生成临时 HTML 文件
        tmp_dir = tempfile.mkdtemp(prefix='onticards_html_')
        tmp_html = os.path.join(tmp_dir, 'report.html')
        with open(tmp_html, 'w', encoding='utf-8') as f:
            f.write(_build_report_html(report, results))

        print(f"[LibreOffice-Exporter] 临时 HTML 文件生成完成: {tmp_html}")

        last_err = None

        # ===== Tier 1: 本地 soffice（优先） =====
        local_soffice_path = _get_local_soffice_path()
        if local_soffice_path:
            print(f"[LibreOffice-Exporter] [Tier 1] 本地 soffice 可用 ({local_soffice_path})，开始转换 {target_fmt}...")
            try:
                result = self._convert_via_local_soffice(tmp_html, target_fmt, file_name)
                if result:
                    print(f"[LibreOffice-Exporter] [Tier 1] ✅ soffice 转换成功 | mode={result['mode']}")
                    self._cleanup_tmp_dir(tmp_dir)
                    return result
            except Exception as e:
                last_err = e
                print(f"[LibreOffice-Exporter] [Tier 1] ❌ soffice 转换失败，降级: {e}")
        else:
            print(f"[LibreOffice-Exporter] [Tier 1] soffice 未找到，跳过")

        # ===== Tier 2: 纯 Python 降级 =====
        print(f"[LibreOffice-Exporter] [Tier 2] 进入纯 Python 降级方案...")
        if target_fmt == 'docx':
            file_path = os.path.join(self.export_path, file_name).replace('\\', '/')
            print(f"[LibreOffice-Exporter] [Tier 2] 使用 python-docx 生成 docx | path={file_path}")
            _build_docx_with_python(report, results, file_path)
            self._cleanup_tmp_dir(tmp_dir)
            return {
                'file_path': file_path,
                'file_name': file_name,
                'file_size': os.path.getsize(file_path),
                'format': 'docx',
                'mode': 'python-docx',
            }
        elif target_fmt == 'pdf':
            # soffice Tier 1 已在上面失败，不再有降级方案
            # 直接抛出错误（PDF 格式暂时不支持）
            self._cleanup_tmp_dir(tmp_dir)
            raise RuntimeError(
                f"PDF 导出失败（soffice 不可用），当前版本暂不支持 PDF 格式。"
            )
        else:
            # 其他格式兜底
            self._cleanup_tmp_dir(tmp_dir)
            raise RuntimeError(
                f"不支持的目标格式: {target_fmt}，可选值: docx, pdf, xlsx"
            )

        self._cleanup_tmp_dir(tmp_dir)
        raise RuntimeError(
            f"文档生成失败（已尝试：soffice），last_error={last_err}"
        )

    def _export_docx(self,
                     report: GovernanceReport,
                     results: List[RuleExecutionResult],
                     custom_file_name: str = None) -> Dict[str, Any]:
        """导出 Word 文档（soffice 优先，python-docx 降级）"""
        return self._build_and_convert_html(report, results, 'docx', custom_file_name=custom_file_name)

    def _export_pdf(self,
                    report: GovernanceReport,
                    results: List[RuleExecutionResult],
                    custom_file_name: str = None) -> Dict[str, Any]:
        """导出 PDF 文档（soffice 优先，weasyprint 降级）"""
        return self._build_and_convert_html(report, results, 'pdf', custom_file_name=custom_file_name)

    def _export_xlsx(self,
                     report: GovernanceReport,
                     results: List[RuleExecutionResult],
                     custom_file_name: str = None) -> Dict[str, Any]:
        """导出 Excel 文档

        生成优先级：
        1. 本地 soffice（subprocess）— 效果最好，优先使用
        2. openpyxl（纯 Python 降级）
        """
        print(f"[LibreOffice-Exporter] 开始生成 Excel | report_id={report.id}")
        file_name = self._make_filename(report, 'xlsx', custom_file_name=custom_file_name)
        tmp_dir = tempfile.mkdtemp(prefix='onticards_xlsx_')

        try:
            # 生成 HTML（Excel 内容）
            tmp_html = os.path.join(tmp_dir, 'report.html')
            with open(tmp_html, 'w', encoding='utf-8') as f:
                f.write(_build_report_html(report, results))

            print(f"[LibreOffice-Exporter] XLSX 临时 HTML 生成完成: {tmp_html}")
            last_err = None

            # ===== Tier 1: 本地 soffice =====
            local_soffice_path = _get_local_soffice_path()
            if local_soffice_path:
                print(f"[LibreOffice-Exporter] [Tier 1-XLSX] 本地 soffice 可用，开始转换 xlsx...")
                try:
                    result = self._convert_via_local_soffice(tmp_html, 'xlsx', file_name)
                    if result:
                        print(f"[LibreOffice-Exporter] [Tier 1-XLSX] ✅ soffice 转换成功")
                        self._cleanup_tmp_dir(tmp_dir)
                        return result
                except Exception as e:
                    last_err = e
                    print(f"[LibreOffice-Exporter] [Tier 1-XLSX] ❌ soffice 转换失败，降级: {e}")
            else:
                print(f"[LibreOffice-Exporter] [Tier 1-XLSX] soffice 未找到，跳过")

            # ===== Tier 2: openpyxl =====
            print(f"[LibreOffice-Exporter] [Tier 2-XLSX] 进入纯 Python 降级方案（openpyxl）...")
            file_path = os.path.join(self.export_path, file_name).replace('\\', '/')
            print(f"[LibreOffice-Exporter] [Tier 2-XLSX] 使用 openpyxl 生成 xlsx | path={file_path}")
            _build_xlsx_with_python(report, results, file_path)
            self._cleanup_tmp_dir(tmp_dir)
            return {
                'file_path': file_path,
                'file_name': file_name,
                'file_size': os.path.getsize(file_path),
                'format': 'xlsx',
                'mode': 'openpyxl',
            }
        finally:
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass
