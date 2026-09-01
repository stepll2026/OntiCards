# -*- coding:utf-8 -*-
"""
@File: test_report_libreoffice.py
@Description: soffice 本地调用测试脚本（完全独立，不依赖 Flask / .env）
@Author: 韩小豪 849631113@qq.com
@Create: 2026-07-06

运行方式（任意目录下执行）：
    python controllers/governance/test_report_libreoffice.py

前置条件：
    pip install requests

测试内容：
    - 检查 soffice 命令是否可用
    - HTML → DOCX 转换
    - HTML → PDF 转换
    - HTML → XLSX 转换
    - python-docx 降级（DOCX）
    - openpyxl 降级（XLSX）

已移除：
    - weasyprint PDF 降级（依赖 GTK，配置复杂且效果不如 soffice）
"""

import os
import sys
import io
import shutil
import tempfile
import subprocess
import re
from datetime import datetime
from typing import Tuple, Dict, Any, Optional, List

# Windows PowerShell UTF-8 输出兼容
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

EXPORT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
os.makedirs(EXPORT_PATH, exist_ok=True)


# ==================== soffice 工具函数 ====================

def _get_local_soffice_path() -> Optional[str]:
    """查找本地 soffice 可执行文件路径（优先 soffice.exe，过滤 .com）"""
    candidates = []
    try:
        if os.name == 'nt':
            result = subprocess.run(
                ['where', 'soffice'],
                capture_output=True, text=True, timeout=5
            )
            if result.returncode == 0:
                candidates = result.stdout.strip().split('\n')
        else:
            result = subprocess.run(
                ['which', 'soffice'],
                capture_output=True, text=True, timeout=5
            )
            if result.returncode == 0:
                candidates = result.stdout.strip().split('\n')
    except Exception:
        pass

    # Windows: 优先 soffice.exe，过滤 .com（.com 对 filter 系统支持不完整）
    if os.name == 'nt':
        exe_paths = [p for p in candidates if p.lower().endswith('.exe') and not p.lower().endswith('.com')]
        com_paths = [p for p in candidates if p.lower().endswith('.com')]
        # 优先 exe 列表第一位，其次 com 列表第一位
        if exe_paths:
            return exe_paths[0]
        if com_paths:
            return com_paths[0]
        # 兜底：常见 Windows 安装路径
        for base in [
            r"C:\Program Files\LibreOffice\program",
            r"C:\Program Files (x86)\LibreOffice\program",
            r"H:\LibreOffice\program",
        ]:
            exe = os.path.join(base, "soffice.exe")
            if os.path.isfile(exe):
                return exe
    else:
        # Linux/Docker: 直接取第一个
        if candidates:
            return candidates[0]
    return None


_FILTER_MAP = {
    'docx': 'HTML (StarWriter)',
    'xlsx': 'HTML (StarCalc)',
}


def _run_local_soffice(
    src_path: str,
    dst_dir: str,
    target_format: str,
    timeout: int = 120,
) -> Optional[str]:
    """调用本地 soffice 转换文件

    Args:
        src_path: 源文件路径
        dst_dir: 输出目录
        target_format: 目标格式（docx, pdf, xlsx, odt, ods）
        timeout: 超时时间（秒）

    Returns:
        输出文件路径，失败返回 None
    """
    soffice = _get_local_soffice_path()
    if not soffice:
        return None

    try:
        cmd = [
            soffice,
            '--headless',
            '--convert-to', target_format,
        ]
        infilter = _FILTER_MAP.get(target_format.lower())
        if infilter:
            cmd.append('--infilter=' + infilter)
        cmd.extend(['--outdir', dst_dir, src_path])
        env = os.environ.copy()
        if os.name == 'nt':
            env['HOME'] = tempfile.gettempdir()
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout,
            env=env,
        )
        if result.returncode != 0:
            print(f"      [soffice stderr] {result.stderr.strip()}")
        base_name = os.path.splitext(os.path.basename(src_path))[0]
        output_name = f"{base_name}.{target_format}"
        output_path = os.path.join(dst_dir, output_name)
        if os.path.exists(output_path):
            return output_path
        for f in os.listdir(dst_dir):
            if f.startswith(base_name) and f.endswith(f'.{target_format}'):
                return os.path.join(dst_dir, f)
        return None
    except subprocess.TimeoutExpired:
        print(f"      [soffice 超时] 命令执行超过 {timeout}s")
        return None
    except Exception as e:
        print(f"      [soffice 异常] {e}")
        return None


# ==================== 纯 Python 降级方案 ====================

def _build_docx_with_python(file_path: str, title: str, content_rows: List[List[str]]):
    """使用 python-docx 生成 DOCX（降级方案）"""
    try:
        from docx import Document
    except ImportError:
        print("      [错误] python-docx 未安装，无法生成 DOCX")
        return

    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("生成模式: python-docx（降级方案）")

    doc.add_heading("数据内容", 1)
    table = doc.add_table(rows=1, cols=len(content_rows[0]) if content_rows else 1)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(content_rows[0]):
        hdr_cells[i].text = str(cell_text)
    for row_data in content_rows[1:]:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)

    doc.save(file_path)
    print(f"      [OK] python-docx 生成成功: {file_path}")


def _build_xlsx_with_python(file_path: str, title: str, content_rows: List[List[str]]):
    """使用 openpyxl 生成 XLSX（降级方案）"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        print("      [错误] openpyxl 未安装，无法生成 XLSX")
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "报告数据"

    # 标题行
    ws['A1'] = title
    ws['A1'].font = Font(name='Microsoft YaHei', size=14, bold=True)
    ws['A1'].alignment = Alignment(horizontal='left', vertical='center')
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(content_rows[0]), 1) if content_rows else 1)

    # 生成时间
    ws['A2'] = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ws['A2'].font = Font(name='Microsoft YaHei', size=10, color="666666")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max(len(content_rows[0]), 1) if content_rows else 1)

    # 数据表头
    start_row = 4
    if content_rows:
        header_fill = PatternFill(start_color="3498DB", end_color="3498DB", fill_type="solid")
        for col_idx, header_text in enumerate(content_rows[0], start=1):
            cell = ws.cell(row=start_row, column=col_idx, value=str(header_text))
            cell.font = Font(name='Microsoft YaHei', size=11, bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 数据行
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        for row_idx, row_data in enumerate(content_rows[1:], start=start_row + 1):
            for col_idx, cell_text in enumerate(row_data, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=str(cell_text))
                cell.alignment = Alignment(horizontal='left', vertical='center')
                cell.border = thin_border
                if row_idx % 2 == 0:
                    cell.fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")

        # 表头也加边框
        for col_idx in range(1, len(content_rows[0]) + 1):
            ws.cell(row=start_row, column=col_idx).border = thin_border

    wb.save(file_path)
    print(f"      [OK] openpyxl 生成成功: {file_path}")


# ==================== 测试 HTML 模板 ====================

def build_test_html(
    report_name: str = "LibreOffice 格式转换测试报告",
    content_rows: Optional[List[List[str]]] = None,
) -> str:
    if content_rows is None:
        content_rows = [
            ["规则名称", "执行状态", "符合率", "说明"],
            ["非空值检查", "通过", "98.5%", "字段完整性验证"],
            ["格式校验", "通过", "95.2%", "数据格式规范"],
            ["重复值检测", "警告", "87.3%", "存在少量重复"],
            ["一致性检查", "失败", "62.1%", "存在数据不一致"],
        ]

    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    table_rows_html = ""
    for i, row in enumerate(content_rows):
        bg = "#f8f9fa" if i % 2 == 0 else "#ffffff"
        table_rows_html += "<tr>"
        for cell in row:
            table_rows_html += f"<td>{cell}</td>"
        table_rows_html += "</tr>"

    return f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<title>{report_name}</title>
<style>
body{{font-family:"Microsoft YaHei","SimHei",Arial,sans-serif;margin:40px;color:#333;}}
h1{{color:#2c3e50;border-bottom:3px solid #3498db;padding-bottom:10px;}}
table{{width:100%;border-collapse:collapse;margin-top:20px;}}
th{{background:#3498db;color:white;padding:10px;text-align:left;}}
td{{padding:8px;border-bottom:1px solid #ddd;}}
tr:nth-child(even){{background:#f8f9fa;}}
</style>
</head><body>
<h1>{report_name}</h1>
<p>生成时间: {now}</p>
<table>
<thead><tr><th>规则名称</th><th>执行状态</th><th>符合率</th><th>说明</th></tr></thead>
<tbody>
{table_rows_html}
</tbody>
</table>
</body></html>"""


# ==================== 测试用例 ====================

def test_soffice_availability() -> Optional[str]:
    """检查 soffice 是否可用"""
    print()
    print("=" * 70)
    print("  [1] soffice 可用性检查")
    print("=" * 70)
    soffice_path = _get_local_soffice_path()
    if soffice_path:
        print(f"    结果: [OK] soffice 可用")
        print(f"    路径: {soffice_path}")
        return soffice_path
    else:
        print(f"    结果: [FAIL] soffice 未找到")
        print(f"    请确认 Dockerfile 中已安装 libreoffice，或在本地已安装 LibreOffice")
        return None


def test_soffice_filter_diagnostic(soffice_path: str):
    """诊断 soffice filter 注册状态（查找关键 filter 是否存在）"""
    print()
    print("=" * 70)
    print("  [1.5] soffice filter 注册诊断")
    print("=" * 70)
    key_filters = {
        'MS Word 2007 XML': 'docx',
        'MS Excel 2007 XML': 'xlsx',
        'HTML (StarWiki)': 'html',
        'Portable Document Format': 'pdf',
    }
    try:
        result = subprocess.run(
            [soffice_path, '--headless', '--help'],
            capture_output=True, text=True, timeout=10
        )
        output = result.stdout + result.stderr
        found = 0
        for fname, label in key_filters.items():
            if fname in output or label.upper() in output.upper():
                print(f"    [OK]  {fname:30s} ({label})")
                found += 1
            else:
                print(f"    [??]  {fname:30s} ({label}) - 未在 help 中确认")
        print(f"\n    注: 若 DOCX/XLSX export filter 缺失, 需重建 LibreOffice 配置")
        print(f"        方法: 重装 LibreOffice 或删除用户配置目录后重试")
        print(f"        配置目录: %APPDATA%\\LibreOffice 或 ~/.config/libreoffice")
    except Exception as e:
        print(f"    [WARN] filter 诊断执行失败: {e}")


def test_conversions_soffice(html_content: str, timeout: int = 120) -> Dict[str, Any]:
    """通过 soffice subprocess 测试各格式转换"""
    print()
    print("=" * 70)
    print("  [2] 格式转换测试（soffice subprocess）")
    print("=" * 70)

    tmp_dir = tempfile.mkdtemp(prefix='lo_test_')
    src_html = os.path.join(tmp_dir, 'source.html')
    with open(src_html, 'w', encoding='utf-8') as f:
        f.write(html_content)

    results = {}
    test_cases = [
        ('docx', 'Word 文档 (DOCX)'),
        ('pdf',  'PDF 文档 (PDF)'),
        ('xlsx', 'Excel 电子表格 (XLSX)'),
    ]

    for to_fmt, label in test_cases:
        print(f"\n  --- {label}（soffice）---")
        output_path = _run_local_soffice(src_html, EXPORT_PATH, to_fmt, timeout=timeout)
        if output_path and os.path.exists(output_path):
            size = os.path.getsize(output_path)
            results[to_fmt] = {
                'success': True,
                'path': output_path,
                'size': size,
                'mode': 'soffice',
            }
            print(f"      [OK] soffice 转换成功")
            print(f"           文件: {output_path}")
            print(f"           大小: {size:,} bytes ({size/1024:.1f} KB)")
        else:
            results[to_fmt] = {'success': False, 'error': 'soffice 转换失败'}
            print(f"      [FAIL] soffice 转换失败或输出文件未找到")

    shutil.rmtree(tmp_dir)
    return results


def test_conversions_python_fallback(html_content: str) -> Dict[str, Any]:
    """测试纯 Python 降级方案（当 soffice 不可用时）"""
    print()
    print("=" * 70)
    print("  [3] 纯 Python 降级方案测试")
    print("=" * 70)

    content_rows = [
        ["规则名称", "执行状态", "符合率", "说明"],
        ["非空值检查", "通过", "98.5%", "字段完整性验证"],
        ["格式校验", "通过", "95.2%", "数据格式规范"],
        ["重复值检测", "警告", "87.3%", "存在少量重复"],
        ["一致性检查", "失败", "62.1%", "存在数据不一致"],
    ]

    results = {}

    # python-docx
    print("\n  --- Word 文档 (python-docx) ---")
    docx_path = os.path.join(EXPORT_PATH, f"fallback_docx_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    try:
        _build_docx_with_python(docx_path, "LibreOffice 格式转换测试报告（降级方案）", content_rows)
        size = os.path.getsize(docx_path)
        results['docx'] = {'success': True, 'path': docx_path, 'size': size, 'mode': 'python-docx'}
        print(f"      [OK] python-docx 生成成功 | {size/1024:.1f} KB")
    except Exception as e:
        results['docx'] = {'success': False, 'error': str(e)}
        print(f"      [FAIL] python-docx: {e}")

    # openpyxl
    print("\n  --- Excel 文档 (openpyxl) ---")
    xlsx_path = os.path.join(EXPORT_PATH, f"fallback_xlsx_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    try:
        _build_xlsx_with_python(xlsx_path, "LibreOffice 格式转换测试报告（降级方案）", content_rows)
        size = os.path.getsize(xlsx_path)
        results['xlsx'] = {'success': True, 'path': xlsx_path, 'size': size, 'mode': 'openpyxl'}
        print(f"      [OK] openpyxl 生成成功 | {size/1024:.1f} KB")
    except Exception as e:
        results['xlsx'] = {'success': False, 'error': str(e)}
        print(f"      [FAIL] openpyxl: {e}")

    return results


# ==================== 主入口 ====================

def main():
    print()
    print("#" * 70)
    print("#  LibreOffice soffice 本地调用测试")
    print("#  不依赖 Flask / .env / HTTP 服务")
    print("#" * 70)
    print(f"\n  Python:    {sys.version.split()[0]}")
    print(f"  soffice:   {'可用' if _get_local_soffice_path() else '不可用'}")
    print(f"  导出目录:  {EXPORT_PATH}")

    # [1] soffice 可用性
    soffice_path = test_soffice_availability()

    # [1.5] filter 注册诊断
    if soffice_path:
        test_soffice_filter_diagnostic(soffice_path)

    # [2] soffice 转换测试
    soffice_results = {}
    if soffice_path:
        html = build_test_html()
        soffice_results = test_conversions_soffice(html, timeout=120)
    else:
        print("    跳过（soffice 不可用）")

    # [3] Python 降级方案测试
    python_results = {}
    html = build_test_html()
    python_results = test_conversions_python_fallback(html)

    # ===== 汇总 =====
    print()
    print("=" * 70)
    print("  测试汇总")
    print("=" * 70)

    all_ok = True

    print("\n  【soffice 转换结果】")
    for fmt, res in soffice_results.items():
        label = fmt.upper()
        if res.get('success'):
            sz = res['size'] / 1024
            print(f"    [OK]  {label:6s}  {sz:7.1f} KB  mode=soffice")
        else:
            print(f"    [FAIL] {label:6s}  {res.get('error', '')}")
            all_ok = False

    print("\n  【纯 Python 降级方案结果】")
    for fmt, res in python_results.items():
        label = fmt.upper()
        if res.get('success'):
            sz = res['size'] / 1024
            print(f"    [OK]  {label:6s}  {sz:7.1f} KB  mode={res.get('mode', 'unknown')}")
        else:
            print(f"    [FAIL] {label:6s}  {res.get('error', '')}")
            all_ok = False

    print()
    print("=" * 70)
    if soffice_path and all_ok:
        print("  [PASS] 所有测试通过！soffice 和 Python 降级方案均运行正常。")
    elif soffice_path:
        print("  [PARTIAL] soffice 可用，但部分降级方案失败。")
    else:
        print("  [WARN] soffice 不可用，请确认容器中已正确安装 LibreOffice。")
        print("         纯 Python 降级方案结果见上方汇总。")
    print("=" * 70)
    print(f"\n  报告文件已保存至: {EXPORT_PATH}")


if __name__ == '__main__':
    main()
